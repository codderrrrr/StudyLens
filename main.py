# main.py

import os
import time
import logging
from pathlib import Path
from typing import List, Tuple, Optional

import numpy as np
from PIL import Image

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse
from pydantic import BaseModel

from paddleocr import PaddleOCR
from docx import Document

import spacy

from transformers import pipeline

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_google_genai import ChatGoogleGenerativeAI

from sentence_transformers import CrossEncoder

import os
from dotenv import load_dotenv

load_dotenv()

# =========================================================
# CONFIG
# =========================================================

MAX_LONG_EDGE = 1400
UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

CHUNK_SIZE = 600
CHUNK_OVERLAP = 150

EMBED_MODEL = "BAAI/bge-base-en-v1.5"
RERANK_MODEL = "cross-encoder/ms-marco-MiniLM-L-6-v2"

MMR_K = 8
MMR_FETCH_K = 20
RERANK_TOP_N = 4


ENV_FLAGS = {
    "FLAGS_use_mkldnn": "0",
    "FLAGS_use_onednn": "0",
    "FLAGS_pir_mkldnn_pass": "0",
    "FLAGS_enable_mkldnn": "0",
    "PADDLE_DISABLE_MKLDNN": "1",
    "OMP_NUM_THREADS": "4",
    "MKL_NUM_THREADS": "4",
}


# =========================================================
# FASTAPI APP
# =========================================================

app = FastAPI(
    title="OCR + RAG Backend",
    version="1.0.0"
)


# =========================================================
# STARTUP
# =========================================================

print("\nLoading models...\n")

os.environ.update(ENV_FLAGS)

for name in ("ppocr", "paddlex", "PIL"):
    logging.getLogger(name).setLevel(logging.ERROR)


# OCR
ocr = PaddleOCR(
    lang="en",
    device="cpu",
    enable_mkldnn=False,
    det_db_thresh=0.3,
    det_db_box_thresh=0.5,
    rec_batch_num=6,
    det_limit_side_len=960,
    use_angle_cls=False,
)

# NLP
nlp = spacy.load("en_core_web_sm")

# Grammar correction
grammar_corrector = pipeline(
    "text2text-generation",
    model="vennify/t5-base-grammar-correction"
)

# Embeddings
embeddings = HuggingFaceEmbeddings(
    model_name=EMBED_MODEL,
    model_kwargs={"device": "cpu"},
    encode_kwargs={"normalize_embeddings": True},
)

# Reranker
reranker = CrossEncoder(RERANK_MODEL)

# Gemini LLM
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    temperature=0.2,
    google_api_key=GOOGLE_API_KEY,
)

print("All models loaded successfully.\n")


# =========================================================
# GLOBAL VECTOR DB
# =========================================================

vector_db = None


# =========================================================
# REQUEST MODELS
# =========================================================

class AskRequest(BaseModel):
    question: str


class GenerateQARequest(BaseModel):
    topic: Optional[str] = None


# =========================================================
# OCR FUNCTIONS
# =========================================================

def preprocess(path: str) -> np.ndarray:
    img = Image.open(path).convert("RGB")

    w, h = img.size
    long_edge = max(w, h)

    if long_edge > MAX_LONG_EDGE:
        scale = MAX_LONG_EDGE / long_edge
        img = img.resize(
            (int(w * scale), int(h * scale)),
            Image.Resampling.LANCZOS
        )

    return np.array(img)


def parse_results(results: list) -> Tuple[List[str], List[Tuple]]:
    lines = []
    regions = []

    if not results:
        return lines, regions

    res = results[0]

    if isinstance(res, dict):

        items = zip(
            res.get("rec_texts", []),
            res.get("rec_scores", []),
            res.get("dt_polys", []),
        )

        for txt, score, bbox in items:
            if txt.strip():
                lines.append(txt)
                regions.append((bbox, txt, score))

    elif isinstance(res, list):

        for item in res:
            if len(item) == 2:
                bbox, (txt, score) = item

                if txt.strip():
                    lines.append(txt)
                    regions.append((bbox, txt, score))

    return lines, regions


def save_docx(lines: List[str], filename: str):
    doc = Document()

    doc.add_heading("OCR Extracted Text", level=1)

    for line in lines:
        line = line.strip()

        if line:
            doc.add_paragraph(line)

    doc.save(filename)


def reconstruct_sentences(text: List[str]) -> str:

    joined = " ".join(text)

    doc = nlp(joined)

    sentences = [sent.text.strip() for sent in doc.sents]

    return " ".join(sentences)


def grammar_correction(text: str) -> str:

    result = grammar_corrector(
        f"grammar: {text}",
        max_length=512
    )

    return result[0]["generated_text"]


# =========================================================
# RAG FUNCTIONS
# =========================================================

def build_chunks(text: str):

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        separators=["\n\n", "\n", ". ", "? ", "! ", " "],
        length_function=len,
        add_start_index=True,
    )

    return splitter.create_documents([text])


def build_vectorstore(text: str):

    docs = build_chunks(text)

    db = FAISS.from_documents(docs, embeddings)

    return db


def rerank_docs(query, docs):

    if not docs:
        return docs

    pairs = [(query, doc.page_content) for doc in docs]

    scores = reranker.predict(pairs)

    ranked = sorted(
        zip(docs, scores),
        key=lambda x: x[1],
        reverse=True
    )

    return [doc for doc, _ in ranked[:RERANK_TOP_N]]


# =========================================================
# API ROUTES
# =========================================================

@app.get("/")
def root():
    return {
        "message": "OCR + RAG Backend Running"
    }


# =========================================================
# OCR ENDPOINT
# =========================================================

@app.post("/ocr")
async def run_ocr(file: UploadFile = File(...)):

    global vector_db

    try:

        start = time.perf_counter()

        if not file.filename:
            raise HTTPException(
                status_code=400,
                detail="Invalid filename"
            )

        file_path = os.path.join(
            UPLOAD_DIR,
            file.filename
        )

        with open(file_path, "wb") as f:
            f.write(await file.read())

        img_np = preprocess(file_path)

        results = list(ocr.predict(img_np))

        lines, regions = parse_results(results)

        reconstructed = reconstruct_sentences(lines)

        corrected_text = grammar_correction(reconstructed)

        # Save TXT
        txt_path = os.path.join(
            OUTPUT_DIR,
            "output.txt"
        )

        with open(txt_path, "w", encoding="utf-8") as f:
            f.write(corrected_text)

        # Save DOCX
        docx_path = os.path.join(
            OUTPUT_DIR,
            "output.docx"
        )

        save_docx(
            corrected_text.split(". "),
            docx_path
        )

        # Build Vector DB
        vector_db = build_vectorstore(corrected_text)

        end = time.perf_counter()

        return JSONResponse({
            "status": "success",
            "text": corrected_text,
            "lines_detected": len(lines),
            "processing_time": round(end - start, 2),
            "docx_file": docx_path
        })

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )


# =========================================================
# ASK QUESTION ENDPOINT
# =========================================================

@app.post("/ask")
async def ask_question(request: AskRequest):

    global vector_db

    if vector_db is None:
        raise HTTPException(
            status_code=400,
            detail="Please run OCR first."
        )

    try:

        retriever = vector_db.as_retriever(
            search_type="mmr",
            search_kwargs={
                "k": MMR_K,
                "fetch_k": MMR_FETCH_K,
                "lambda_mult": 0.6,
            },
        )

        raw_docs = retriever.invoke(request.question)

        reranked_docs = rerank_docs(
            request.question,
            raw_docs
        )

        context = "\n\n".join(
            [doc.page_content for doc in reranked_docs]
        )

        prompt = f"""
You are an academic assistant.

CONTEXT:
{context}

QUESTION:
{request.question}

RULES:
- Answer only from context if possible
- If not in context, clearly mention it
- Keep answer concise
"""

        response = llm.invoke(prompt)

        return JSONResponse({
            "question": request.question,
            "answer": response.content,
            "sources": [
                doc.page_content[:200]
                for doc in reranked_docs
            ]
        })

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )


# =========================================================
# GENERATE EXAM QUESTIONS
# =========================================================

@app.post("/generate-questions")
async def generate_questions():

    global vector_db

    if vector_db is None:
        raise HTTPException(
            status_code=400,
            detail="Please run OCR first."
        )

    try:

        retriever = vector_db.as_retriever(
            search_type="mmr",
            search_kwargs={
                "k": 8,
                "fetch_k": 20,
                "lambda_mult": 0.7,
            },
        )

        docs = retriever.invoke(
            "important exam topics"
        )

        context = "\n\n".join(
            [doc.page_content for doc in docs]
        )

        prompt = f"""
You are a university professor.

Generate 6 exam-style questions from the notes below.

NOTES:
{context}

Requirements:
- 2 Easy
- 3 Medium
- 1 Hard
- Include conceptual and application questions
"""

        result = llm.invoke(prompt)

        return JSONResponse({
            "questions": result.content
        })

    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=str(e)
        )


# =========================================================
# RUN
# =========================================================

"""
Run Server:

uvicorn main:app --reload

Swagger Docs:

http://127.0.0.1:8000/docs
"""