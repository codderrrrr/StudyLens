import os
import time
import logging
import numpy as np
from PIL import Image
from paddleocr import PaddleOCR
import spacy
from happytransformer import HappyTextToText, TTSettings
from docx import Document # type: ignore
# ── Constants ─────────────────────────────────────────────────────────────────
MAX_LONG_EDGE = 1400
IMAGE_PATH    = "image.png"

ENV_FLAGS = {
    "FLAGS_use_mkldnn":      "0",
    "FLAGS_use_onednn":      "0",
    "FLAGS_pir_mkldnn_pass": "0",
    "FLAGS_enable_mkldnn":   "0",
    "PADDLE_DISABLE_MKLDNN": "1",
    "OMP_NUM_THREADS":       "4",
    "MKL_NUM_THREADS":       "4",
}


# ── Setup ─────────────────────────────────────────────────────────────────────
def configure_env() -> None:
    os.environ.update(ENV_FLAGS)
    for name in ("ppocr", "paddlex", "PIL"):
        logging.getLogger(name).setLevel(logging.ERROR)


def build_ocr() -> PaddleOCR:
    return PaddleOCR(
        lang="en",
        device="cpu",
        enable_mkldnn=False,
        det_db_thresh=0.3,
        det_db_box_thresh=0.5,
        rec_batch_num=6,
        det_limit_side_len=960,
        use_angle_cls=False,
    )


# ── Image preprocessing ───────────────────────────────────────────────────────
def preprocess(path: str) -> np.ndarray:
    img = Image.open(path).convert("RGB")
    w, h = img.size
    long_edge = max(w, h)
    if long_edge > MAX_LONG_EDGE:
        scale = MAX_LONG_EDGE / long_edge
        img   = img.resize((int(w * scale), int(h * scale)), Image.LANCZOS) # type: ignore
        print(f"  Resized: {w}x{h} → {img.size[0]}x{img.size[1]}")
    else:
        print(f"  Image size: {w}x{h} (no resize needed)")
    return np.array(img)


# ── Result parsing ────────────────────────────────────────────────────────────
def parse_results(results: list) -> tuple[list[str], list[tuple]]:
    """Return (lines, regions) from PaddleOCR predict output."""
    lines, regions = [], []
    if not results:
        return lines, regions

    res = results[0]

    if isinstance(res, dict):                                   # PaddleOCR 3.x
        items = zip(
            res.get("rec_texts",  []),
            res.get("rec_scores", []),
            res.get("dt_polys",   []),
        )
        for txt, score, bbox in items:
            if txt.strip():
                lines.append(txt)
                regions.append((bbox, txt, score))

    elif isinstance(res, list):                                 # legacy API
        for item in res:
            if len(item) == 2:
                bbox, (txt, score) = item
                if txt.strip():
                    lines.append(txt)
                    regions.append((bbox, txt, score))

    return lines, regions


def print_results(lines: list[str], regions: list[tuple]) -> None:
    print("\n===== TEXT (reading order) =====\n")
    print("\n".join(lines))

    print("\n===== LAYOUT (bbox + confidence) =====\n")
    for bbox, txt, score in regions:
        if bbox is not None:
            xs = [p[0] for p in bbox]
            ys = [p[1] for p in bbox]
            x, y = int(min(xs)), int(min(ys))
            w    = int(max(xs) - min(xs))
            h    = int(max(ys) - min(ys))
            print(f"  [{x:4d},{y:4d}] {w:3d}x{h:3d}  {score:.2f}  {txt}")
        else:
            print(f"  [no bbox]  {score:.2f}  {txt}")

def save_to_doc(text : list[str]):
    doc = Document()
    
    doc.add_heading("OCR Extracted Text", level=1)
    
    for line in text:
        line = line.strip()
        if line:  # skip empty lines
            doc.add_paragraph(line)

    doc.save("output.docx")
    
def read_docx(file_path: str):
    doc = Document(file_path)
    
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    
    return text

def sentence_reconstruction(text):
    if isinstance(text, list):
        text = " ".join(text)
    nlp = spacy.load("en_core_web_sm")
    
    doc = nlp(text=text)
    sentences = [sent.text.strip() for sent in doc.sents]
    return " ".join(sentences)
    
def run_ocr(image_path: str = IMAGE_PATH) -> tuple[list[str], list[tuple]]:
    # configure_env()

    print("Loading model...")
    t0  = time.perf_counter()
    ocr = build_ocr()
    print(f"Model loaded in {time.perf_counter() - t0:.1f}s")

    img_np = preprocess(image_path)

    print("Running OCR...")
    t1      = time.perf_counter()
    results = list(ocr.predict(img_np))
    t_ocr   = time.perf_counter() - t1
    print(f"OCR done in {t_ocr:.1f}s")

    lines, regions = parse_results(results)
    print_results(lines, regions)
    
    save_to_doc(lines)

    t_total = time.perf_counter() - t0
    print(f"\n⏱️  Total: {t_total:.1f}s  |  OCR only: {t_ocr:.1f}s")
    return lines, regions

def correct_ocr():
    
    corrector = pipeline( # type: ignore
        "text2text-generation",
        model="vennify/t5-base-grammar-correction"
    )
    
    text = read_docx("output.docx")
    
    sentence_reconstruction(text)
    
    save_to_doc(text)
    print(text)

if __name__ == "__main__":
    # run_ocr('image.png')
    correct_ocr()